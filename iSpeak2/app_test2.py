import os

# Set cache paths
os.environ['HF_HOME'] = '/slowfs/amstpubs1/sanchit/iSpeak/iSpeak_Final/cache/'
os.environ['TRANSFORMERS_CACHE'] = '/slowfs/amstpubs1/sanchit/iSpeak/iSpeak_Final/cache/'

import csv
import zipfile
import pandas as pd
import numpy as np
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify
from gtts import gTTS
import torch
import torchaudio
from speechbrain.pretrained import Tacotron2, HIFIGAN
from transformers import AutoProcessor, BarkModel
import scipy
from nltk.tokenize import sent_tokenize
from docx import Document
from werkzeug.utils import secure_filename
import random
import json
from model import NeuralNet as TTSNeuralNet
from nltk_utils import bag_of_words, tokenize
from flask_cors import CORS
import inflect
from edgetts import edgetts_bp
from edgetts import text_to_speech  # Import Edge TTS generation function

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Change this to a random secret key
CORS(app)

# Register the TTS blueprint
app.register_blueprint(edgetts_bp)

# Create an inflect engine
p = inflect.engine()

def convert_numbers_to_words(text):
   # Split text into tokens
   tokens = text.split()
   # Convert each token if it's a number
   converted_tokens = [p.number_to_words(token) if token.isdigit() else token for token in tokens]
   # Join tokens back to a single string
   return ' '.join(converted_tokens)

# Device configuration: 
cpu_device = torch.device("cpu")
gpu_device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# Load pre-trained Tacotron2 model (using CPU)
tacotron2 = Tacotron2.from_hparams(source="speechbrain/tts-tacotron2-ljspeech", savedir="tmpdir_tts").to(cpu_device)

# Load pre-trained HiFIGAN vocoder model (using CPU)
hifi_gan = HIFIGAN.from_hparams(source="speechbrain/tts-hifigan-ljspeech", savedir="tmpdir_vocoder").to(cpu_device)

# Load Bark processor and model (using GPU)
processor = AutoProcessor.from_pretrained("suno/bark")
model = BarkModel.from_pretrained("suno/bark").to(gpu_device)


# Path to the CSV files
csv_file = 'replacement.csv'
users_file = 'users.csv'
replacement_df = pd.read_csv("replacement_words.csv")

# Paths for the new feature
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
CSV_FILE = 'replacements_for_transcript.csv'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['CSV_FILE'] = CSV_FILE

# Ensure the folders exist
if not os.path.exists(UPLOAD_FOLDER):
   os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(PROCESSED_FOLDER):
   os.makedirs(PROCESSED_FOLDER)

# Function to update the CSV file with new values
def update_csv(original_text, replacement_text):
   with open(csv_file, 'a', newline='') as file:
       writer = csv.writer(file)
       writer.writerow([original_text, replacement_text])

# Function to check if user is logged in
def login_required(f):
   def wrap(*args, **kwargs):
       if 'logged_in' in session:
           return f(*args, **kwargs)
       else:
           return redirect(url_for('login'))
   wrap.__name__ = f.__name__
   return wrap

@app.route('/')
@login_required
def index():
   return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
   if request.method == 'POST':
       username = request.form['username']
       password = request.form['password']
       
       with open(users_file, 'r') as file:
           reader = csv.DictReader(file)
           for row in reader:
               if row['username'] == username and row['password'] == password:
                   session['logged_in'] = True
                   return redirect(url_for('index'))
       return "Login failed. Check your username and password."
   return render_template('login.html')

@app.route('/logout', methods=['POST'])
def logout():
   session.pop('logged_in', None)
   return redirect(url_for('login'))

@app.route('/generate_audio', methods=['POST'])
@login_required
def generate_audio():
   text = request.form['text']
   voice = request.form['voice']
   preset = request.form.get('preset')

   # Replace words in the text with their replacements based on the CSV file
   for index, row in replacement_df.iterrows():
       text = text.replace(row['original_text'], row['replacement_text'])

   if voice == 'tacotron2':
       audio_path = generate_audio_tacotron2(text)
   elif voice == 'gtts':
       audio_path = generate_audio_gtts(text)
   elif voice == 'bark_suno':
       audio_path = generate_audio_bark(text, preset)
   return render_template('preview.html', audio_path=audio_path)

@app.route('/download_audio', methods=['POST'])
@login_required
def download_audio():
   audio_path = request.form['audio_path']
   return send_file(audio_path, as_attachment=True)

def generate_unique_filename(extension):
   return f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"


def generate_audio_tacotron2(text):
   # Convert numbers to words
   text = convert_numbers_to_words(text)

   with torch.no_grad():
       mel_output, _, _ = tacotron2.encode_text(text)
       mel_output = mel_output.cpu()

   with torch.no_grad():
       waveforms = hifi_gan.decode_batch(mel_output)
       waveforms = waveforms.squeeze(1).cpu().numpy()

   audio_path = f"static/{generate_unique_filename('wav')}"
   torchaudio.save(audio_path, torch.tensor(waveforms), 22050)
   return audio_path

def generate_audio_gtts(text):
   tts = gTTS(text=text, lang='en')
   audio_path = f"static{generate_unique_filename('mp3')}"
   tts.save(audio_path)
   return audio_path

def generate_audio_bark(text, preset):
   sentences = sent_tokenize(text)
   audio_list = []

   for sentence in sentences:
       inputs = processor(sentence, voice_preset=preset)
       for k, v in inputs.items():
           inputs[k] = v.to(gpu_device)
       audio_array = model.generate(**inputs)
       audio_array = audio_array.cpu().numpy().squeeze()
       audio_list.append(audio_array)

       silence_duration = int(model.generation_config.sample_rate * 0.25)
       silence = np.zeros(silence_duration)
       audio_list.append(silence)

   concatenated_audio = np.concatenate(audio_list)
   stereo_audio = np.stack([concatenated_audio, concatenated_audio], axis=1)

   audio_path = f"static/{generate_unique_filename('wav')}"
   if not os.path.exists('static'):
       os.makedirs('static')
   scipy.io.wavfile.write(audio_path, rate=model.generation_config.sample_rate, data=stereo_audio)
   return audio_path
   
@app.route('/form')
@login_required
def form():
   return render_template('form.html')

@app.route('/submit_form', methods=['POST'])
@login_required
def submit_form():
   original_text = request.form['original_text']
   replacement_text = request.form['replacement_text']
   update_csv(original_text, replacement_text)
   return redirect(url_for('index'))

@app.route('/bulk_audio', methods=['GET', 'POST'])
@login_required
def bulk_audio():
   if request.method == 'POST':
       doc_file = request.files['doc_file']
       voice = request.form['voice']
       preset = request.form.get('preset')

       # Save the uploaded DOC file
       doc_path = f"uploads/{doc_file.filename}"
       doc_file.save(doc_path)

       # Extract paragraphs from the DOC file
       paragraphs = extract_paragraphs_from_doc(doc_path)

       # Generate audio for each paragraph
       audio_files = []
       for i, paragraph in enumerate(paragraphs):
           if voice == 'tacotron2':
               audio_path = generate_audio_tacotron2(paragraph)
           elif voice == 'gtts':
               audio_path = generate_audio_gtts(paragraph)
           elif voice == 'bark_suno':
               audio_path = generate_audio_bark(paragraph, preset)
           elif voice == 'edge_tts':  # New Edge TTS option
               audio_path = text_to_speech(paragraph)
           audio_files.append(audio_path)

       # Create a zip file containing all audio files
       zip_path = create_zip_file(audio_files)

       return send_file(zip_path, as_attachment=True)

   return render_template('bulk_audio.html')

def extract_paragraphs_from_doc(doc_path):
   doc = Document(doc_path)
   paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
   return paragraphs

def create_zip_file(file_paths):
   zip_filename = f"audio_files_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
   zip_path = f"static/{zip_filename}"
   with zipfile.ZipFile(zip_path, 'w') as zipf:
       for file in file_paths:
           zipf.write(file, os.path.basename(file))
   return zip_path

# New feature from App2
@app.route('/transcript', methods=['GET', 'POST'])
@login_required
def transcript():
   if request.method == 'POST':
       if 'file' not in request.files:
           return redirect(request.url)
       file = request.files['file']
       if file.filename == '':
           return redirect(request.url)
       if file:
           filename = secure_filename(file.filename)
           input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
           file.save(input_path)

           output_filename = f"processed_{filename}"
           output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
           process_transcript(input_path, app.config['CSV_FILE'], output_path)

           return redirect(url_for('download_file', filename=output_filename))
   return render_template('transcript_processor.html')

@app.route('/download/<filename>')
@login_required
def download_file(filename):
   return send_file(os.path.join(app.config['PROCESSED_FOLDER'], filename), as_attachment=True)

def process_transcript(input_file, csv_file, output_file):
   # Read the CSV file and create
   replacements = {}
   with open(csv_file, newline='') as csvfile:
       reader = csv.reader(csvfile)
       for row in reader:
           if len(row) == 2:
               replacements[row[0]] = row[1]

   doc = Document(input_file)
   paragraphs = []
   current_note = []

   for para in doc.paragraphs:
       if para.text.startswith("Slide"):
           if current_note:
               paragraphs.append(" ".join(current_note))
               current_note = []
           slide_note = para.text.split(" ", 2)
           if len(slide_note) > 2:
               current_note.append(slide_note[2])
       else:
           current_note.append(para.text)

   if current_note:
       paragraphs.append(" ".join(current_note))

   new_doc = Document()
   for note in paragraphs:
       for word, replacement in replacements.items():
           note = note.replace(word, replacement)
       new_doc.add_paragraph(note)

   new_doc.save(output_file)

@app.route('/feature_request')
@login_required
def feature_request():
   return render_template('feature_request.html')

@app.route('/submit_feature_request', methods=['POST'])
@login_required
def submit_feature_request():
   name = request.form['name']
   email = request.form['email']
   feature = request.form['feature']
   priority = request.form['priority']
   feature_request_file = 'feature_requests.csv'

   with open(feature_request_file, 'a', newline='') as file:
       writer = csv.writer(file)
       writer.writerow([name, email, feature, priority, datetime.now().strftime('%Y-%m-%d %H:%M:%S')])

   return redirect(url_for('index'))

#Chatbot
# Load chatbot model, vectorizer, and label encoder
chatbot_device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

with open('intents.json', 'r') as json_data:
   chatbot_intents = json.load(json_data)

chatbot_file = "data.pth"
chatbot_data = torch.load(chatbot_file)

chatbot_input_size = chatbot_data["input_size"]
chatbot_hidden_size = chatbot_data["hidden_size"]
chatbot_output_size = chatbot_data["output_size"]
chatbot_all_words = chatbot_data['all_words']
chatbot_tags = chatbot_data['tags']
chatbot_model_state = chatbot_data["model_state"]

chatbot_model = TTSNeuralNet(chatbot_input_size, chatbot_hidden_size, chatbot_output_size).to(chatbot_device)
chatbot_model.load_state_dict(chatbot_model_state)
chatbot_model.eval()

def chatbot_get_response(msg):
   sentence = tokenize(msg)
   X = bag_of_words(sentence, chatbot_all_words)
   X = X.reshape(1, X.shape[0])
   X = torch.from_numpy(X).to(chatbot_device).float()  # Convert to torch.FloatTensor

   output = chatbot_model(X)
   _, predicted = torch.max(output, dim=1)

   tag = chatbot_tags[predicted.item()]

   probs = torch.softmax(output, dim=1)
   prob = probs[0][predicted.item()]
   if prob.item() > 0.75:
       for intent in chatbot_intents['intents']:
           if tag == intent["tag"]:
               return random.choice(intent['responses'])
   else:
       return "I do not understand..."

@app.route("/")
@login_required
def home():
   return render_template("index.html")

@app.route("/api/get_response", methods=["POST"])
@login_required
def api_get_response():
   user_input = request.json.get("message")
   response = chatbot_get_response(user_input)
   return jsonify({"response": response})


if __name__ == '__main__':
   if not os.path.exists('uploads'):
       os.makedirs('uploads')
   if not os.path.exists('processed'):
       os.makedirs('processed')
   app.run(host='0.0.0.0', port=5000, debug=False)