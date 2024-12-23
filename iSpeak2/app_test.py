import os
import csv
import pandas as pd
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, session, send_file
from werkzeug.utils import secure_filename
from flask_cors import CORS
from edgetts import edgetts_bp
from speaker_notes_bp import speaker_notes_bp
import csv

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Change this to a random secret key
CORS(app)

# Register the TTS blueprint
app.register_blueprint(edgetts_bp)

# Register the new feature Blueprint
app.register_blueprint(speaker_notes_bp, url_prefix='/speaker_notes')

# Path to the CSV files
csv_file = 'replacement.csv'
users_file = 'users.csv'
replacement_df = pd.read_csv("replacement_words.csv")

# Ensure folders exist
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
CSV_FILE = 'replacements_for_transcript.csv'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['CSV_FILE'] = CSV_FILE


if not os.path.exists(UPLOAD_FOLDER):
   os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(PROCESSED_FOLDER):
   os.makedirs(PROCESSED_FOLDER)


# Function to update the CSV file with new values
def update_csv(original_text, replacement_text):
   with open(csv_file, 'a', newline='') as file:
       writer = csv.writer(file)
       writer.writerow([original_text, replacement_text])

# Function to check if user is logged in
def login_required(f):
   def wrap(*args, **kwargs):
       if 'logged_in' in session:
           return f(*args, **kwargs)
       else:
           return redirect(url_for('login'))
   wrap.__name__ = f.__name__
   return wrap


@app.route('/')
@login_required
def index():
   return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
   if request.method == 'POST':
       username = request.form['username']
       password = request.form['password']
       with open(users_file, 'r') as file:
           reader = csv.DictReader(file)
           for row in reader:
               if row['username'] == username and row['password'] == password:
                   session['logged_in'] = True
                   return redirect(url_for('index'))
       return "Login failed. Check your username and password."
   return render_template('login.html')

@app.route('/logout', methods=['POST'])
def logout():
   session.pop('logged_in', None)
   return redirect(url_for('login'))

@app.route('/form')
@login_required
def form():
   return render_template('form.html')

@app.route('/submit_form', methods=['POST'])
def submit_form():
   original_text = request.form['original_text']
   replacement_text = request.form['replacement_text']
   update_csv(original_text, replacement_text)
   return redirect(url_for('index'))

@app.route('/transcript', methods=['GET', 'POST'])
@login_required
def transcript():
   if request.method == 'POST':
       if 'file' not in request.files:
           return redirect(request.url)
       file = request.files['file']
       if file.filename == '':
           return redirect(request.url)
       if file:
           filename = secure_filename(file.filename)
           input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
           file.save(input_path)

           output_filename = f"processed_{filename}"
           output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
           process_transcript(input_path, app.config['CSV_FILE'], output_path)

           return redirect(url_for('download_file', filename=output_filename))
   return render_template('transcript_processor.html')

@app.route('/download/<filename>')
def download_file(filename):
   return send_file(os.path.join(app.config['PROCESSED_FOLDER'], filename), as_attachment=True)

def process_transcript(input_file, csv_file, output_file):
   # Read the CSV file and create replacements
   replacements = {}
   with open(csv_file, newline='') as csvfile:
       reader = csv.reader(csvfile)
       for row in reader:
           if len(row) == 2:
               replacements[row[0]] = row[1]

   from docx import Document
   doc = Document(input_file)
   paragraphs = []
   current_note = []

   for para in doc.paragraphs:
       if para.text.startswith("Slide"):
           if current_note:
               paragraphs.append(" ".join(current_note))
               current_note = []
           slide_note = para.text.split(" ", 2)
           if len(slide_note) > 2:
               current_note.append(slide_note[2])
       else:
           current_note.append(para.text)

   if current_note:
       paragraphs.append(" ".join(current_note))

   new_doc = Document()
   for note in paragraphs:
       for word, replacement in replacements.items():
           note = note.replace(word, replacement)
       new_doc.add_paragraph(note)

   new_doc.save(output_file)

if __name__ == '__main__':
   app.run(host='0.0.0.0', port=5000, debug=False)